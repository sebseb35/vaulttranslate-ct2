from pathlib import Path

import pytest
from docx import Document

from packages.adapter_docx import DocxDocumentAdapter
from packages.core.models import DocumentFormat, Segment


FIXTURES_DIR = Path(__file__).resolve().parents[1] / "fixtures" / "docx"


def _fixture_bytes(name: str) -> bytes:
    return (FIXTURES_DIR / name).read_bytes()


def test_docx_adapter_extracts_paragraphs_runs_tables_header_footer() -> None:
    adapter = DocxDocumentAdapter()
    source = _fixture_bytes("sample.docx")

    segments = adapter.extract_segments(source)
    texts = [segment.text for segment in segments]

    assert adapter.supported_format is DocumentFormat.DOCX
    assert all(segment.segment_id.startswith("docx-") for segment in segments)
    assert "Header text" in texts
    assert "Hello " in texts
    assert "world" in texts
    assert "Second paragraph" in texts
    assert "Cell A1" in texts
    assert "Cell B2" in texts
    assert "Footer text" in texts


def test_docx_adapter_extracts_hyperlinks_and_mixed_inline_content() -> None:
    adapter = DocxDocumentAdapter()
    source = _fixture_bytes("fidelity_hardening.docx")

    segments = adapter.extract_segments(source)
    texts = [segment.text for segment in segments]

    assert "Prefix " in texts
    assert "OpenAI Docs" in texts
    assert " suffix" in texts
    assert "Line 1\nLine 2\tTabbed" in texts
    assert "Only Link" in texts
    assert "Cell prefix " in texts
    assert "Cell Link" in texts
    assert "Header baseline" in texts


def test_docx_adapter_rebuild_roundtrip_preserves_structure_and_updates_text(tmp_path: Path) -> None:
    adapter = DocxDocumentAdapter()
    source = _fixture_bytes("sample.docx")
    extracted = adapter.extract_segments(source)

    translated = tuple(
        Segment(segment_id=segment.segment_id, text=f"{segment.text} FR")
        for segment in extracted
    )
    rebuilt_bytes = adapter.rebuild_document(source, translated)

    output_path = tmp_path / "rebuilt.docx"
    output_path.write_bytes(rebuilt_bytes)
    doc = Document(output_path)

    assert doc.paragraphs[0].runs[0].text == "Hello  FR"
    assert doc.paragraphs[0].runs[1].text == "world FR"
    assert doc.paragraphs[0].runs[1].bold is True
    assert doc.paragraphs[1].text == "Second paragraph FR"
    assert doc.tables[0].cell(0, 0).text == "Cell A1 FR"
    assert doc.tables[0].cell(1, 1).text == "Cell B2 FR"
    assert doc.sections[0].header.paragraphs[0].text == "Header text FR"
    assert doc.sections[0].footer.paragraphs[0].text == "Footer text FR"


def test_docx_adapter_rebuild_preserves_hyperlinks_and_inline_formatting(tmp_path: Path) -> None:
    adapter = DocxDocumentAdapter()
    source = _fixture_bytes("fidelity_hardening.docx")
    extracted = adapter.extract_segments(source)

    translated = tuple(
        Segment(segment_id=segment.segment_id, text=f"{segment.text} FR")
        for segment in extracted
    )
    rebuilt_bytes = adapter.rebuild_document(source, translated)

    output_path = tmp_path / "rebuilt-fidelity.docx"
    output_path.write_bytes(rebuilt_bytes)
    doc = Document(output_path)

    first_para = doc.paragraphs[0]
    mixed_content = list(first_para.iter_inner_content())
    assert mixed_content[0].text == "Prefix  FR"
    assert mixed_content[1].text == "OpenAI Docs FR"
    assert mixed_content[1].url == "https://example.com/docs"
    assert mixed_content[2].text == " suffix FR"

    second_para = doc.paragraphs[1]
    assert second_para.runs[0].text == "Line 1\nLine 2\tTabbed FR"
    assert second_para.runs[0].italic is True

    link_only_para = doc.paragraphs[2]
    only_link_content = list(link_only_para.iter_inner_content())
    assert only_link_content[1].text == "Only Link FR"
    assert only_link_content[1].url == "https://example.com/only"

    cell_content = list(doc.tables[0].cell(0, 0).paragraphs[0].iter_inner_content())
    assert cell_content[0].text == "Cell prefix  FR"
    assert cell_content[1].text == "Cell Link FR"
    assert cell_content[1].url == "https://example.com/cell"
    assert doc.sections[0].header.paragraphs[0].text == "Header baseline FR"


def test_docx_adapter_rebuild_validates_segment_count_and_ids() -> None:
    adapter = DocxDocumentAdapter()
    source = _fixture_bytes("sample.docx")
    extracted = adapter.extract_segments(source)

    with pytest.raises(ValueError, match="same number of segments"):
        adapter.rebuild_document(source, extracted[:-1])

    wrong_first = (Segment(segment_id="docx-9999", text="bad"),) + extracted[1:]
    with pytest.raises(ValueError, match="segment mismatch"):
        adapter.rebuild_document(source, wrong_first)
