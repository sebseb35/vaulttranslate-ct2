from pathlib import Path

from docx import Document
from pypdf import PdfReader
from typer.testing import CliRunner

from apps.cli.main import app
from packages.core.models import Segment, TranslationResult, TranslationRequest
from tests.unit.test_pdf_adapter import FIXTURES_DIR

runner = CliRunner()


def test_help_shows_translate_command() -> None:
    result = runner.invoke(app, ["--help"])

    assert result.exit_code == 0
    assert "translate" in result.stdout


def test_translate_dry_run_outputs_summary(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--dry-run",
        ],
    )

    assert result.exit_code == 0
    assert "Dry-run: no translation executed." in result.output
    assert "document_format=txt" in result.output
    assert f"input={input_path}" in result.output
    assert f"output={output_path}" in result.output
    assert "segments=1" in result.output
    assert "engine=mock" in result.output


def test_translate_rejects_unsupported_extension(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.odt"
    output_path = tmp_path / "out.odt"
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--dry-run",
        ],
    )

    assert result.exit_code != 0
    assert "Unsupported input format" in result.output


def test_translate_rejects_missing_output_directory(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "missing" / "out.txt"
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--dry-run",
        ],
    )

    assert result.exit_code != 0
    assert "Output directory does not exist" in result.output


def test_translate_writes_output_file_for_txt(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert "Wrote translated file to" in result.output
    assert output_path.exists()
    assert output_path.read_text(encoding="utf-8") == "hello [fr]"


def test_translate_writes_output_file_for_markdown(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.md"
    output_path = tmp_path / "out.md"
    input_path.write_text(
        "Text with `code` and [link](https://example.com).\n\n```python\nprint('hi')\n```\n",
        encoding="utf-8",
    )

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()
    output = output_path.read_text(encoding="utf-8")
    assert "Text with `code` and [link](https://example.com). [fr]" in output
    assert "```python" in output
    assert "print('hi')" in output


def test_translate_writes_output_file_for_docx(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.docx"
    output_path = tmp_path / "out.docx"

    doc = Document()
    p = doc.add_paragraph()
    p.add_run("Hello ")
    bold_run = p.add_run("world")
    bold_run.bold = True
    doc.sections[0].header.paragraphs[0].text = "Header text"
    doc.save(input_path)

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()

    translated_doc = Document(output_path)
    assert translated_doc.paragraphs[0].runs[0].text == "Hello "
    assert translated_doc.paragraphs[0].runs[1].text == "world [fr]"
    assert translated_doc.paragraphs[0].runs[1].bold is True
    assert translated_doc.sections[0].header.paragraphs[0].text == "Header text [fr]"


def test_translate_writes_output_file_for_pdf(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.pdf"
    output_path = tmp_path / "out.pdf"
    input_path.write_bytes((FIXTURES_DIR / "sample.pdf").read_bytes())

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()
    reader = PdfReader(str(output_path))
    rebuilt_text = "\n".join(page.extract_text() or "" for page in reader.pages)
    assert "Quarterly report [fr]" in rebuilt_text
    assert "Revenue up 12% [fr]" in rebuilt_text


def test_translate_uses_ct2_mode_when_model_path_is_provided(
    tmp_path: Path,
    monkeypatch,
) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    model_dir = tmp_path / "ct2-model"
    tokenizer_dir = tmp_path / "tokenizer"
    model_dir.mkdir()
    tokenizer_dir.mkdir()
    input_path.write_text("hello", encoding="utf-8")

    captured: dict[str, object] = {}

    class _FakeEngine:
        def translate(self, request: TranslationRequest) -> TranslationResult:
            translated_segments = tuple(
                Segment(segment_id=s.segment_id, text=f"{s.text} [REAL]", constraints=s.constraints)
                for s in request.segments
            )
            return TranslationResult(
                document_format=request.document_format,
                source_language=request.source_language,
                target_language=request.target_language,
                translated_segments=translated_segments,
                metadata={**request.metadata, "engine": "fake-ct2"},
            )

    def _fake_select_engine(
        *,
        model_path: Path | None,
        tokenizer_path: str | None,
        inter_threads: int,
        intra_threads: int,
        compute_type: str,
    ) -> _FakeEngine:
        captured["model_path"] = model_path
        captured["tokenizer_path"] = tokenizer_path
        captured["inter_threads"] = inter_threads
        captured["intra_threads"] = intra_threads
        captured["compute_type"] = compute_type
        return _FakeEngine()

    monkeypatch.setattr("apps.cli.main.select_engine", _fake_select_engine)

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--model-path",
            str(model_dir),
            "--tokenizer-path",
            str(tokenizer_dir),
            "--inter-threads",
            "2",
            "--intra-threads",
            "3",
            "--compute-type",
            "int8",
        ],
    )

    assert result.exit_code == 0
    assert "engine=ctranslate2" in result.output
    assert f"model_path={model_dir}" in result.output
    assert output_path.read_text(encoding="utf-8") == "hello [REAL]"
    assert captured == {
        "model_path": model_dir,
        "tokenizer_path": str(tokenizer_dir),
        "inter_threads": 2,
        "intra_threads": 3,
        "compute_type": "int8",
    }


def test_translate_rejects_tokenizer_path_without_model_path(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--tokenizer-path",
            str(tmp_path / "tokenizer"),
        ],
    )

    assert result.exit_code != 0
    assert "--tokenizer-path requires --model-path" in result.output


def test_translate_rejects_missing_tokenizer_directory(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    model_dir = tmp_path / "ct2-model"
    missing_tokenizer_dir = tmp_path / "missing-tokenizer"
    model_dir.mkdir()
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--model-path",
            str(model_dir),
            "--tokenizer-path",
            str(missing_tokenizer_dir),
        ],
    )

    assert result.exit_code != 0
    assert "Tokenizer path does not exist" in result.output


def test_translate_rejects_model_path_when_not_directory(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    model_file = tmp_path / "model.bin"
    model_file.write_text("x", encoding="utf-8")
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--model-path",
            str(model_file),
        ],
    )

    assert result.exit_code != 0
    assert "Model path must be a directory" in result.output


def test_translate_rejects_tokenizer_path_when_not_directory(tmp_path: Path) -> None:
    input_path = tmp_path / "sample.txt"
    output_path = tmp_path / "out.txt"
    model_dir = tmp_path / "ct2-model"
    tokenizer_file = tmp_path / "tokenizer.json"
    model_dir.mkdir()
    tokenizer_file.write_text("{}", encoding="utf-8")
    input_path.write_text("hello", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
            "--model-path",
            str(model_dir),
            "--tokenizer-path",
            str(tokenizer_file),
        ],
    )

    assert result.exit_code != 0
    assert "Tokenizer path must be a directory" in result.output
