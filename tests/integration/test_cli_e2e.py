from __future__ import annotations

from pathlib import Path

from docx import Document
from typer.testing import CliRunner

from apps.cli.main import app

runner = CliRunner()


def test_cli_e2e_txt_writes_output(tmp_path: Path) -> None:
    input_path = tmp_path / "input.txt"
    output_path = tmp_path / "output.txt"
    input_path.write_text("Alpha line.\nBeta line.\n", encoding="utf-8")

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()
    assert output_path.read_text(encoding="utf-8") == "Alpha line.\nBeta line. [fr]\n"


def test_cli_e2e_md_writes_output(tmp_path: Path) -> None:
    input_path = tmp_path / "input.md"
    output_path = tmp_path / "output.md"
    input_path.write_text(
        "Hello markdown.\n\n`code` and [link](https://example.com)\n\n```python\nprint('x')\n```\n",
        encoding="utf-8",
    )

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()
    output = output_path.read_text(encoding="utf-8")
    assert "Hello markdown. [fr]" in output
    assert "`code` and [link](https://example.com) [fr]" in output
    assert "```python" in output
    assert "print('x')" in output


def test_cli_e2e_docx_writes_output(tmp_path: Path) -> None:
    fixture_path = Path(__file__).resolve().parents[1] / "fixtures" / "docx" / "sample.docx"
    input_path = tmp_path / "input.docx"
    output_path = tmp_path / "output.docx"
    input_path.write_bytes(fixture_path.read_bytes())

    result = runner.invoke(
        app,
        [
            "translate",
            "--input",
            str(input_path),
            "--output",
            str(output_path),
            "--source",
            "en",
            "--target",
            "fr",
        ],
    )

    assert result.exit_code == 0
    assert output_path.exists()

    translated_doc = Document(output_path)
    assert translated_doc.paragraphs[0].runs[0].text == "Hello "
    assert translated_doc.paragraphs[0].runs[1].text == "world [fr]"
    assert translated_doc.paragraphs[1].text == "Second paragraph [fr]"
    assert translated_doc.tables[0].cell(0, 0).text == "Cell A1 [fr]"
    assert translated_doc.sections[0].header.paragraphs[0].text == "Header text [fr]"
    assert translated_doc.sections[0].footer.paragraphs[0].text == "Footer text [fr]"
